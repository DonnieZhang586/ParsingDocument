 #!/usr/bin/env python

# This script is for python 3.x
# This script is to parse input documents and extract content from them.
# The result will be written to a txt file in json format.
# The script can be modified to insert the data into ElasticSearch.

# How to use it:
#  Make sure the script and the folder are under the same directory
#  Make sure you are in the right working directory
#  In terminal, print:
#      python ScriptName FolderName

'''
# **INSTALLATION GUIDE**
## **pdfminer**
pip install pdfminer.six

## **wand**
pip install Wand
### _Wand dependencies_
#### mac OS
_install homebrew first_
brew install imagemagick@6 --with-liblqr
_In bash shell, please type:_
ln -s /usr/local/Cellar/imagemagick@6/<your specific 6 version>/lib/libMagickWand-6.Q16.dylib /usr/local/lib/libMagickWand.dylib
#### linux ubuntu
_wand does not get along with imagemagick 7.x, please check:_
https://askubuntu.com/questions/936583/install-imagemagick-for-python-ubuntu-16-04
http://docs.wand-py.org/en/0.4.4/guide/install.html
#### windows
_Please check:_
http://docs.wand-py.org/en/0.4.4/guide/install.html

## **Image**
pip install Pillow
_For more information please go to:_
 https://pillow.readthedocs.io/en/5.1.x/installation.html

## **pyocr**
### Python 2.7
sudo pip install pyocr
### Python 3.x
sudo pip3 install pyocr

## **docx**
pip install python-docx

## **textract**
### Ubuntu/Debian
apt-get install python-dev libxml2-dev libxslt1-dev antiword unrtf poppler-utils pstotext tesseract-ocr \
flac ffmpeg lame libmad0 libsox-fmt-mp3 sox libjpeg-dev swig
pip install textract
### mac OS
_These steps rely on you having homebrew installed as well as the cask plugin (brew install caskroom/cask/brew-cask)_
brew cask install xquartz
brew install poppler antiword unrtf tesseract swig
pip install textract
'''




import os.path
import os
import sys
import json

from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.converter import TextConverter
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfpage import PDFTextExtractionNotAllowed
from pdfminer.pdfinterp import PDFResourceManager
from pdfminer.pdfinterp import PDFPageInterpreter
from pdfminer.layout import LAParams
from pdfminer.converter import PDFPageAggregator
import pdfminer
from io import StringIO


# for mac OSX please download ImageMagick@6, because wand does not get along with ImageMagick7.x
from wand.image import Image
from PIL import Image as PI
import pyocr
import pyocr.builders
import io

#python 3.x pip install python-docx
from docx import Document 

import textract 


ext = ['.doc','.pptx','.odt','.ppt','.xls','.xlsx','.txt','.html']
pic = ['.jpeg', '.jpg', '.png', '.gif', '.bmp']

def CLSFC_file(file):
    #categorize the file type
    if file.endswith(".pdf"):
        return Parse_PDF(file)
    elif file.endswith(".docx"):
        return Parse_Docx(file)
    elif file.endswith(tuple(ext)):
        return Parse_Doc(file)
    elif file.endswith(tuple(pic)):
        return Parse_Pic(file)
    else:
        print("WARNING - Could not parse document " +str(file))


def Parse_PDF(FileName):
    try:
        # Open and read the pdf file
        fp = open(FileName, 'rb')
        # Create parser object to parse the pdf content
        parser = PDFParser(fp)
        # No password for the pdf file
        password = ""
        document = PDFDocument(parser, password)
        # check out password protection
        if not document.is_extractable:
            print("File Under Password Protection"+str(FileName))
            raise PDFTextExtractionNotAllowed("File Under Password Protection")
        # Create PDFResourceManager object that stores shared resources such as fonts or images
        rsrcmgr = PDFResourceManager()
        retstr = StringIO()
        codec = 'utf-8'
        laparams = LAParams()
        # Create a PDFDevice object which translates interpreted information into desired format
        # Device needs to be connected to resource manager to store shared resources
        device = TextConverter(rsrcmgr, retstr, codec=codec, laparams=laparams)
        interpreter = PDFPageInterpreter(rsrcmgr, device)        
        maxpages = 0
        caching = True
        pagenos=set()
        for page in PDFPage.get_pages(fp, 
                                      pagenos, 
                                      maxpages=maxpages, 
                                      password=password,
                                      caching=caching):
            interpreter.process_page(page)
        fp.close()
        device.close()
        output = retstr.getvalue()
        retstr.close()
        
        # If content is extracted successfully, return the content, otherwise go to pyOCR
        if len(output) != 1:
            return output
        else:
            # go to pyOCR
            # Initialize pyOCR
            # The tools are returned in the recommended order of usage
            tool = pyocr.get_available_tools()[0]
            req_image = []
            txt = ""
            image_pdf = Image(filename = FileName, resolution = 300)
            image_jpeg = image_pdf.convert('jpeg')
            for img in image_jpeg.sequence:
                img_page = Image(image = img)
                req_image.append(img_page.make_blob('jpeg'))
            for img in req_image:
                tem= tool.image_to_string( PI.open(io.BytesIO(img)),
                                           builder = pyocr.builders.TextBuilder())
                txt = txt + tem
            return txt
    except:
        print("ERROR - Could not parse file" + str(file))


def Parse_Docx(DocxName):
    document = Document(DocxName)
    outputDocx = ""
    for para in document.paragraphs:
        outputDocx = outputDocx + para.text + " "
    return outputDocx


def Parse_Doc(DocName):
    text = textract.process(DocName)
    return text.decode()


def Parse_Pic(PicName):
    tool = pyocr.get_available_tools()[0]
    langs = tool.get_available_languages()
    lang = langs[0]
    plainstr = tool.image_to_string(
        PI.open(PicName),
        lang=lang,
        builder=pyocr.builders.TextBuilder()
        )
    return plainstr



home = ".."
IN_FILE_PATH = os.path.join(home, sys.argv[1])

MyList = []
for dirpath, dirs, files in os.walk(IN_FILE_PATH,topdown = True):
    for file in files:
        FILE_DETAIL_PATH = os.path.join(dirpath,file)
        a = CLSFC_file(FILE_DETAIL_PATH)
        print("Parsing File" + str(file))

        if a:
            a = a.encode("utf-8","ignore")
            MyDictionary = {}
            MyDictionary['FileName'] = file
            FILE_DIRECTORY = os.path.dirname(FILE_DETAIL_PATH)
            MyDictionary['Directory'] = FILE_DIRECTORY
            a = a.decode("ascii", "ignore")
            MyDictionary['Content'] = a.replace("\n"," ").replace("\t"," ").replace("\f"," ")
            MyList.append(MyDictionary)
with open('data8.txt', 'a+') as outfile:
    json.dump(MyList, outfile)

   